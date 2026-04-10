"""
监督智能体配套工具 - preview 解析、数据校验、文本处理、报告生成

将 supervisor_agent 中的纯函数封装到此模块，便于复用与维护。
"""
import ast
import logging
import os
import re
from typing import Dict, List, Optional, Tuple

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, Cm

from tools.data_loader import load_dataframe, resolve_column

logger = logging.getLogger(__name__)


# ---------- 路径与数据链 ----------


def extract_data_files(work_dir: str) -> List[Dict[str, str]]:
    """
    从工作目录提取可下载的数据文件（清洗、归一化、岩性、储层等 CSV）。
    返回 [{"filename": "xxx_cleaned.csv", "label": "清洗后数据"}, ...]
    """
    result: List[Dict[str, str]] = []
    if not work_dir or not os.path.isdir(work_dir):
        return result
    labels = (
        ("_cleaned.csv", "清洗后数据"),
        ("_minmax_normalized.csv", "归一化数据"),
        ("_standard_normalized.csv", "归一化数据"),
        ("_robust_normalized.csv", "归一化数据"),
        ("_zscore_normalized.csv", "归一化数据"),
        ("_normalized.csv", "归一化数据"),
        ("_lithology.csv", "岩性解释结果"),
        ("_reservoir.csv", "储层识别结果"),
        ("_mud_gas_zones.csv", "录井气测高值段汇总"),
        ("_mud_gas_metrics.csv", "录井气测逐点指标"),
    )
    for f in sorted(os.listdir(work_dir)):
        if not f.endswith(".csv"):
            continue
        for suffix, label in labels:
            if f.endswith(suffix):
                result.append({"filename": f, "label": label})
                break
    return result


def compute_tool_output_path(input_path: str, tool_name: str, params: Optional[Dict] = None) -> str:
    """根据工具和输入路径计算输出路径（用于数据链传递）。clean/normalize 统一输出 CSV。"""
    if not input_path or not os.path.exists(input_path):
        return input_path
    dir_name = os.path.dirname(input_path)
    base_name = os.path.basename(input_path)
    name, _ = os.path.splitext(base_name)
    if tool_name == "clean_data":
        return os.path.join(dir_name, f"{name}_cleaned.csv")
    if tool_name == "normalize_data":
        method = (params or {}).get("method", "minmax")
        return os.path.join(dir_name, f"{name}_{method}_normalized.csv")
    if tool_name == "interpret_lithology":
        return os.path.join(dir_name, f"{name}_lithology.csv")
    if tool_name == "identify_reservoir":
        return os.path.join(dir_name, f"{name}_reservoir.csv")
    if tool_name == "analyze_mud_gas_survey":
        return os.path.join(dir_name, f"{name}_mud_gas_zones.csv")
    return input_path


# ---------- Preview 解析 ----------

def extract_columns_from_preview(preview_result: str) -> str:
    """从 preview 数据预览文本结果中提取列名信息，供任务规划使用。"""
    if not preview_result:
        return "（未能获取列信息）"
    lines = preview_result.split("\n")
    cols_raw = ""
    for i, line in enumerate(lines):
        if "【列名】" in line or "列名" in line:
            if "：" in line or ":" in line:
                part = line.split("：")[-1].split(":")[-1].strip()
                if part and len(part) > 2:
                    cols_raw = part
                    break
            if i + 1 < len(lines):
                cols_raw = lines[i + 1].strip()
                break
    if not cols_raw and "列名" in preview_result:
        match = re.search(r"列名[：:]\s*(.+)", preview_result)
        if match:
            cols_raw = match.group(1).strip()
    if not cols_raw:
        return preview_result[:500] + "..." if len(preview_result) > 500 else preview_result
    cols_list = [c.strip() for c in cols_raw.replace("，", ",").split(",") if c.strip()]
    cols_str = ", ".join(cols_list)
    return f"""可用列名（parameters 中必须原样使用以下名称）: {cols_str}
提示：深度列常见为 depth 或 Depth，请根据实际列名填写 depth_column。"""


def get_columns_list_from_preview(preview_result: str) -> List[str]:
    """从 preview 结果中提取列名列表"""
    if not preview_result:
        return []
    lines = preview_result.split("\n")
    for i, line in enumerate(lines):
        if "【列名】" in line or "列名" in line:
            if "：" in line or ":" in line:
                part = line.split("：")[-1].split(":")[-1].strip()
                if part and len(part) > 2:
                    return [c.strip() for c in part.replace("，", ",").split(",") if c.strip()]
            if i + 1 < len(lines):
                raw = lines[i + 1].strip()
                return [c.strip() for c in raw.replace("，", ",").split(",") if c.strip()]
    match = re.search(r"列名[：:]\s*(.+)", preview_result)
    if match:
        raw = match.group(1).strip()
        return [c.strip() for c in raw.replace("，", ",").split(",") if c.strip()]
    return []


def parse_preview_metadata(preview_result: str) -> tuple:
    """
    从 preview 文本中解析列类型与样本数据，用于结构推断。
    返回 (col_dtypes: dict, col_samples: dict)，解析失败返回 ({}, {})。
    """
    col_dtypes, col_samples = {}, {}
    if not preview_result:
        return col_dtypes, col_samples
    dtype_match = re.search(r"【数据类型】\s*\n?\s*(\{[^{}]+\})", preview_result)
    if dtype_match:
        try:
            raw = dtype_match.group(1).strip()
            col_dtypes = ast.literal_eval(raw)
        except Exception:
            pass
    sample_start = preview_result.find("【前")
    if sample_start >= 0:
        tail = preview_result[sample_start:]
        brace_start = tail.find("{")
        if brace_start >= 0:
            raw = tail[brace_start:]
            depth, i, n = 0, 0, len(raw)
            for i in range(n):
                if raw[i] == "{":
                    depth += 1
                elif raw[i] == "}":
                    depth -= 1
                    if depth == 0:
                        break
            try:
                full = ast.literal_eval(raw[: i + 1])
                for col, idx_map in (full or {}).items():
                    if isinstance(idx_map, dict):
                        vals = [idx_map.get(k) for k in sorted(idx_map.keys())]
                        col_samples[col] = vals
            except Exception:
                pass
    return col_dtypes, col_samples


# ---------- 数据校验 ----------

def _is_numeric_str(s: str) -> bool:
    """判断字符串是否可解析为数值。"""
    try:
        float(s.replace(",", ""))
        return True
    except (ValueError, TypeError):
        return False


def _is_monotonic_numeric(values: list) -> bool:
    """判断数值序列是否单调（升序或降序），用于识别深度列。"""
    nums = []
    for v in values:
        if v is None or (isinstance(v, float) and (pd.isna(v) or v != v)):
            continue
        try:
            n = float(v) if not isinstance(v, (int, float)) else v
            nums.append(n)
        except (ValueError, TypeError):
            return False
    if len(nums) < 2:
        return True
    diffs = [nums[i + 1] - nums[i] for i in range(len(nums) - 1)]
    all_pos = all(d >= -1e-9 for d in diffs)
    all_neg = all(d <= 1e-9 for d in diffs)
    return all_pos or all_neg


def _normalize_col_name(name: str) -> str:
    """列名标准化：统一大小写并去除常见分隔符，便于别名匹配。"""
    if not name:
        return ""
    s = str(name).strip().lower()
    # 去掉空格、下划线、连字符、斜杠、括号等常见符号
    return re.sub(r"[\s_\-()/\\]+", "", s)


def _contains_any_alias(col_name: str, aliases: set) -> bool:
    """判断标准化列名是否命中某类曲线别名。"""
    norm = _normalize_col_name(col_name)
    if not norm:
        return False
    if norm in aliases:
        return True
    # 允许部分包含（如 depth_m、grapi、deepresistivity 等）
    for a in aliases:
        if a and (a in norm or norm in a):
            return True
    return False


def _column_has_resistivity(col_name: str) -> bool:
    """
    识别电阻率曲线列：RT/ILD/LLD 及 R10、R20、R90、Rx 等编号命名。
    """
    norm = _normalize_col_name(col_name)
    if not norm:
        return False
    res_aliases = {
        "rt", "resistivity", "电阻率", "ild", "lld", "lls", "rdeep", "rshallow",
        "deepresistivity", "shallowresistivity", "rd", "rs", "rxo", "rxoohm",
    }
    if _contains_any_alias(col_name, res_aliases):
        return True
    if re.match(r"^r\d{1,4}$", norm):
        return True
    if norm.startswith("rx") and len(norm) <= 5:
        return True
    return False


def validate_well_log_data(preview_result: str, cols_list: list) -> tuple:
    """
    基于数据结构智能校验是否为测井/地质数据。
    返回 (True, None) 表示通过；(False, 提示信息) 表示非测井数据。
    """
    if not cols_list or len(cols_list) < 2:
        return False, "数据列数过少，请确认是否为测井数据文件。"

    # 必备测井参数（支持中英文与常见缩写）
    required_groups = [
        ("深度", {
            "depth", "深度", "md", "tvd", "井深", "测深", "depthm", "depthft",
        }),
        ("伽马", {
            "gr", "gammaray", "gamma", "自然伽马", "伽马", "自然伽玛", "cgr",
        }),
        ("电阻率", {
            "rt", "resistivity", "电阻率", "ild", "lld", "lls", "rdeep", "rshallow",
            "deepresistivity", "shallowresistivity", "rd", "rs", "rxo", "rxoohm",
        }),
        ("中子", {
            "cnl", "nphi", "neutron", "中子", "中子孔隙度", "neu", "phin", "cn",
        }),
        ("密度", {
            "den", "rhob", "density", "密度", "体积密度", "bulkdensity",
        }),
        ("声波时差", {
            "dt", "ac", "sonic", "deltat", "声波", "声波时差", "时差", "dtco",
        }),
    ]

    missing_groups = []
    for group_name, aliases in required_groups:
        if group_name == "电阻率":
            has_group = any(_column_has_resistivity(c) for c in cols_list)
        else:
            has_group = any(_contains_any_alias(c, aliases) for c in cols_list)
        if not has_group:
            missing_groups.append(group_name)

    if missing_groups:
        return False, (
            "当前文件缺少测井解释必需参数，无法进行岩性解释/储层识别。"
            f"\n缺失项：{', '.join(missing_groups)}"
            "\n必需参数：深度、伽马(GR)、电阻率(RT/LLD/LLS 等)、中子(CNL/NPHI)、密度(DEN/RHOB)、声波时差(DT/AC)。"
            "\n请上传包含以上关键曲线的 CSV/Excel。"
        )

    col_dtypes, col_samples = parse_preview_metadata(preview_result)

    numeric_cols = set()
    dtypes_map = col_dtypes or {}
    for col in cols_list:
        if not col:
            continue
        dtype_str = str(dtypes_map.get(col, dtypes_map.get(col.strip(), ""))).lower()
        if "float" in dtype_str or "int" in dtype_str:
            numeric_cols.add(col)

    if not numeric_cols and col_samples:
        for col, vals in col_samples.items():
            if vals:
                valid = [v for v in vals[:10] if v is not None and not (isinstance(v, float) and pd.isna(v))]
                if valid and all(isinstance(v, (int, float)) for v in valid):
                    numeric_cols.add(col)

    depth_col = None
    for col in cols_list:
        if not col:
            continue
        cl = col.strip().lower()
        if "depth" in cl or "深度" in cl:
            depth_col = col
            numeric_cols.add(col)
            break
    if not depth_col and col_samples:
        for col in cols_list:
            vals = col_samples.get(col, [])
            if len(vals) >= 2 and _is_monotonic_numeric(vals):
                depth_col = col
                numeric_cols.add(col)
                break

    n_numeric = len(numeric_cols)
    n_total = len(cols_list)
    numeric_ratio = n_numeric / n_total if n_total else 0

    if n_numeric < 2:
        return False, (
            "数据中数值列过少（测井数据通常有多条曲线如 GR、DEN、CNL 等均为数值）。"
            "请确认上传的是包含深度及若干测井曲线的 CSV 或 Excel 文件。"
        )

    if not depth_col:
        return False, (
            "未检测到深度列（测井数据通常有一列随井深单调变化的数值）。"
            "请确认数据中包含深度列（列名可为 depth、Depth、深度 等），或确保至少有一列数值随深度单调递变。"
        )

    if numeric_ratio < 0.4:
        return False, (
            "数据中数值列占比过低（测井数据以数值曲线为主）。"
            "请确认上传的是测井数据文件，而非人员、订单等业务表格。"
        )

    return True, None


def validate_mud_logging_data(cols_list: list) -> tuple:
    """
    校验是否为录井气测类表（深度 + 钻时 + 气测组分/全烃）。
    返回 (True, None) 通过；(False, 提示信息) 不通过。
    """
    if not cols_list or len(cols_list) < 3:
        return False, "录井气测数据列数过少（至少需要深度、钻时与气测相关列）。"

    depth_aliases = {
        "depth", "深度", "md", "tvd", "井深", "测深", "dept", "mdft",
    }
    rop_aliases = {"rop", "钻时", "机械钻速", "钻速", "drillingtime"}
    has_depth = any(_contains_any_alias(c, depth_aliases) for c in cols_list)
    has_rop = any(_contains_any_alias(c, rop_aliases) for c in cols_list)
    has_gas = False
    for c in cols_list:
        nc = _normalize_col_name(c)
        if nc in ("tg", "c1", "c2", "c3", "co2", "other", "全烃", "全量", "甲烷"):
            has_gas = True
            break
        if _contains_any_alias(
            c,
            {"tg", "甲烷", "全烃", "全量", "气测", "乙烷", "丙烷", "丁烷", "戊烷", "二氧化碳", "非烃", "烃"},
        ):
            has_gas = True
            break
        uc = str(c).strip().upper()
        if re.match(r"^(C[1-5]|IC4|NC4|IC5|NC5|TG|CO2|OTHER)$", uc):
            has_gas = True
            break

    if not has_depth:
        return False, "未识别到深度列（井深、Depth、MD 等）。录井气测表需要深度索引。"
    if not has_gas:
        return False, "未识别到气测相关列（如 Tg/全烃、C1 甲烷、C2…C5 等）。"
    if not has_rop:
        return False, "未识别到钻时列（Rop、钻时）。录井分析建议同时包含钻时与气测数据。"
    return True, None


def validate_plot_task_feasibility(
    tool_name: str, input_path: str, params: Optional[Dict] = None
) -> tuple:
    """
    绘图前可行性校验：若数据无法支持该图表，则跳过执行并返回原因。
    返回 (can_run: bool, skip_reason: str)。can_run=True 时 skip_reason 为空。
    """
    plot_tools = {
        "plot_well_log_curves",
        "plot_lithology_distribution",
        "plot_crossplot",
        "plot_heatmap",
        "plot_reservoir_profile",
        "plot_mud_gas_profile",
    }
    if tool_name not in plot_tools:
        return True, ""

    params = params or {}
    try:
        df = load_dataframe(input_path)
    except Exception as e:
        return False, f"无法读取数据文件: {e}"

    cols = list(df.columns)
    numeric_cols = [c for c in cols if c in df.select_dtypes(include=["number"]).columns]

    def _find_depth():
        for c in cols:
            cl = (c or "").lower()
            if "depth" in cl or "深度" in cl:
                return c
        return resolve_column(df, "depth") or resolve_column(df, "Depth") or ""

    depth_col = _find_depth()

    if tool_name == "plot_well_log_curves":
        if not depth_col:
            return False, "数据中未找到深度列，无法绘制测井曲线图"
        if len(numeric_cols) <= 1 or (len(numeric_cols) == 1 and depth_col in numeric_cols):
            return False, "除深度外无其他数值曲线，无法绘制测井曲线图"
        return True, ""

    if tool_name == "plot_lithology_distribution":
        lith_col = resolve_column(df, params.get("lithology_column", "Lithology"))
        if not lith_col:
            for cand in ("Lithology", "lithology", "岩性"):
                lith_col = resolve_column(df, cand)
                if lith_col:
                    break
        if not lith_col:
            return False, "数据中未找到岩性列，无法绘制岩性分布图（需先执行岩性解释）"
        return True, ""

    if tool_name == "plot_crossplot":
        x_param = params.get("x_parameter") or params.get("x") or ""
        y_param = params.get("y_parameter") or params.get("y") or ""
        if not x_param or not y_param:
            return False, "未指定交会图的 x 或 y 参数"
        x_col = resolve_column(df, x_param) or (x_param if x_param in cols else "")
        y_col = resolve_column(df, y_param) or (y_param if y_param in cols else "")
        if not x_col:
            return False, f"未找到 x 轴参数列 '{x_param}'，可用列: {', '.join(cols[:15])}{'...' if len(cols) > 15 else ''}"
        if not y_col:
            return False, f"未找到 y 轴参数列 '{y_param}'，可用列: {', '.join(cols[:15])}{'...' if len(cols) > 15 else ''}"
        if x_col not in numeric_cols or y_col not in numeric_cols:
            return False, "交会图要求 x、y 参数均为数值列"
        return True, ""

    if tool_name == "plot_heatmap":
        if len(numeric_cols) < 2:
            return False, "相关性热力图需要至少 2 个数值列，当前数据不满足"
        return True, ""

    if tool_name == "plot_reservoir_profile":
        if not depth_col:
            return False, "数据中未找到深度列，无法绘制储层剖面图"
        gr_col = resolve_column(df, "GR") or resolve_column(df, "gr")
        poro_col = resolve_column(df, "Porosity") or resolve_column(df, "PHIT") or resolve_column(df, "porosity")
        if not gr_col and not poro_col and len(numeric_cols) <= 1:
            return False, "数据中既无 GR 也无孔隙度等曲线，无法绘制储层剖面图"
        return True, ""

    if tool_name == "plot_mud_gas_profile":
        if not depth_col:
            return False, "数据中未找到深度列，无法绘制录井气测剖面"
        non_depth_numeric = [c for c in numeric_cols if c != depth_col]
        if len(non_depth_numeric) < 1:
            return False, "除深度外无钻时/气测数值列，无法绘制录井气测剖面"
        return True, ""

    return True, ""


# ---------- 文本处理 ----------

def strip_task_json(text: str) -> str:
    """清理模型输出中的内部 JSON 代码块（任务规划等），避免展示到报告/界面。"""
    if not text:
        return text
    text = re.sub(r"```[\s\S]*?```", "", text, flags=re.IGNORECASE)
    stripped = text.lstrip()
    if stripped.startswith("{"):
        depth, end = 0, -1
        for i, c in enumerate(stripped):
            if c == "{":
                depth += 1
            elif c == "}":
                depth -= 1
                if depth == 0:
                    end = i
                    break
        if end >= 0:
            rest = stripped[end + 1 :].lstrip()
            text = strip_task_json(rest) if rest.startswith("{") else rest
    return text.strip()


def streamable_text_for_report(buffer: str) -> Tuple[str, str]:
    """
    流式生成报告时：从累积缓冲中取出「可立即推送到前端」的正文，去掉开头的任务规划 JSON、```json``` 围栏等。

    若开头为未闭合的围栏或未闭合的 `{...}`，返回 ("", buffer)，不向用户展示。

    返回 (emit, carry)：emit 为本次应推送的片段；carry 为尚未展示的后缀缓冲。
    """
    if not buffer:
        return "", ""
    b = buffer
    # 1) 去掉开头的完整 markdown 代码围栏（常为 ```json ... ```）
    while True:
        stripped = b.lstrip()
        if not stripped.startswith("```"):
            break
        nl = stripped.find("\n", 3)
        if nl < 0:
            return "", buffer
        rest_after_lang = stripped[nl + 1 :]
        close = rest_after_lang.find("```")
        if close < 0:
            return "", buffer
        b = b[: len(b) - len(stripped)] + rest_after_lang[close + 3 :]

    stripped = b.lstrip()
    ws_lead = b[: len(b) - len(stripped)]
    if not stripped.startswith("{"):
        return b, ""
    depth = 0
    for i, c in enumerate(stripped):
        if c == "{":
            depth += 1
        elif c == "}":
            depth -= 1
            if depth == 0:
                rest = ws_lead + stripped[i + 1 :]
                return streamable_text_for_report(rest)
    return "", buffer


# ---------- 报告生成 ----------

def extract_report_only(text: str) -> str:
    """从 LLM 输出中提取仅报告部分，去除对话式前导。保留从 # 开头标题开始的内容。"""
    if not text or not text.strip():
        return text
    text = text.strip()
    lines = text.split("\n")
    start_idx = 0
    for i, line in enumerate(lines):
        stripped = line.strip()
        if stripped.startswith("# ") or stripped.startswith("## ") or stripped.startswith("### "):
            start_idx = i
            break
    return "\n".join(lines[start_idx:]).strip() if start_idx > 0 else text


def _strip_asterisks_for_docx(text: str) -> str:
    """移除文本中的星号，避免 docx 报告中出现 * 符号。"""
    if not text:
        return text
    return text.replace("*", "").strip()


def build_interpretation_report_docx(
    work_dir: str, report_content: str, file_path: str
) -> Optional[str]:
    """
    将 Markdown 报告内容写入 Word 文档，并嵌入同目录下的 PNG 图表。
    返回报告相对路径（如 /static/{task_id}/interpretation_report.docx），失败返回 None。
    """
    try:
        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)
        normal_style = doc.styles["Normal"]
        normal_style.font.name = "宋体"
        normal_style.font.size = Pt(10.5)
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing = 1.25
        for level in range(1, 4):
            style_name = f"Heading {level}"
            if style_name in doc.styles:
                h = doc.styles[style_name]
                h.font.name = "黑体"
                h.font.size = Pt(16 - level * 2)
                h.font.bold = True
                h.paragraph_format.space_before = Pt(12)
                h.paragraph_format.space_after = Pt(6)

        clean = _strip_asterisks_for_docx
        for line in report_content.splitlines():
            stripped = line.strip()
            if stripped.startswith("### "):
                doc.add_heading(clean(stripped[4:]), level=3)
            elif stripped.startswith("## "):
                doc.add_heading(clean(stripped[3:]), level=2)
            elif stripped.startswith("# "):
                doc.add_heading(clean(stripped[2:]), level=1)
            elif stripped.startswith("- "):
                p = doc.add_paragraph(clean(stripped[2:]))
                p.style = "List Bullet"
            else:
                doc.add_paragraph(clean(stripped))

        chart_files = [
            f for f in os.listdir(work_dir)
            if (
                f.endswith((".png", ".html"))
                and f != os.path.basename(file_path)
                and not f.startswith(".")
            )
        ]
        # 按「同名无后缀」分组：解释报告优先嵌入静态 PNG（与 Plotly 导出的 HTML 成对出现）
        stems: Dict[str, List[str]] = {}
        for f in chart_files:
            stem, _ = os.path.splitext(f)
            stems.setdefault(stem, []).append(f)

        if stems:
            doc.add_heading("图表", level=1)
            for stem in sorted(stems.keys()):
                try:
                    files = stems[stem]
                    png_f = next((x for x in files if x.lower().endswith(".png")), None)
                    html_f = next((x for x in files if x.lower().endswith(".html")), None)
                    title = stem.replace("_", " ").strip()
                    doc.add_paragraph(title, style="Normal")
                    if png_f:
                        chart_path = os.path.join(work_dir, png_f)
                        if os.path.isfile(chart_path):
                            doc.add_picture(chart_path, width=Inches(5.5))
                            doc.add_paragraph()
                    elif html_f:
                        doc.add_paragraph(
                            "（本图仅有交互式 HTML，请在 Web 界面「可视化」区域查看；静态 PNG 未生成时可检查绘图日志。）",
                            style="Normal",
                        )
                        doc.add_paragraph()
                except Exception:
                    pass

        doc.save(os.path.join(work_dir, "interpretation_report.docx"))
        task_id = os.path.basename(work_dir)
        return f"/static/{task_id}/interpretation_report.docx"
    except Exception:
        return None
